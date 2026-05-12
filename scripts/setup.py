"""First-run scaffolding: create master.md, template-config.yaml, and a
placeholder template.docx that demonstrates the slot syntax.

Refuses to overwrite existing files.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parent.parent

MASTER_MD = PROJECT_ROOT / "master.md"
TEMPLATE_CONFIG = PROJECT_ROOT / "template-config.yaml"
TEMPLATE_DOCX = PROJECT_ROOT / "template.docx"


MASTER_MD_SKELETON = """# Master Profile

This is the single source of truth for all resume content. The writer agent
sources every bullet and every skill from this file. Do not invent claims —
if a metric or technology isn't here, it won't appear in any tailored resume.

Each role/project gets an H2 section. Use dense prose; the writer extracts
bullets per JD. Include numbers, scope, technologies, and outcomes.

---

## Role: <Title> @ <Company> (<Start> – <End or Present>)

Dense prose describing what you did. Include team size, budget, scope,
specific technologies, concrete metrics (numbers, percentages, scale),
business outcomes, and any awards/recognition. The writer will extract
3–5 bullets from this depending on `template-config.yaml`.

Example:
"Built the order-routing service from scratch in Go, handling ~500 RPS
in production. Owned on-call rotation for the payments team (8 engineers),
reduced P1 incidents by 40% through alerting cleanup and runbook authoring.
Led the migration of 4 services to k8s on EKS. Mentored 2 junior engineers
through promotion. Tech: Go, k8s, GitHub Actions, Postgres, Redis, Kafka."

---

## Role: <Previous Title> @ <Previous Company> (<Start> – <End>)

Same shape.

---

## Project: <Project Name>

Dense prose for projects (open source, side, freelance, hackathon).
Include stars/users/revenue/scale, the problem it solves, and the stack.

---

## Skills Inventory

Three categories. The writer picks/orders within each per JD relevance.

### Languages
Python, TypeScript, Go, SQL, Bash

### AI&ML
PyTorch, LangChain, RAG, embeddings, fine-tuning, prompt engineering

### Concepts & Tools
REST APIs, distributed systems, microservices, Docker, k8s, GitHub Actions,
Postgres, Redis, AWS (Lambda, S3, RDS), Terraform
"""


TEMPLATE_CONFIG_SKELETON = """# Slot definitions for template.docx.
# Each id must match a Jinja slot in template.docx (e.g. id "acme_swe"
# corresponds to the slot {%p for b in acme_swe_bullets %}{{b}}{%p endfor %}).

roles:
  - id: example_role_1
    bullet_count: 4
  - id: example_role_2
    bullet_count: 3

projects:
  - id: example_project_1
    bullet_count: 3

skills:
  categories:
    - Languages
    - AI&ML
    - Concepts & Tools
  max_per_category: 8
"""


def add_bullet_loop(doc: Document, slot_id: str) -> None:
    """Add the three-paragraph for/content/endfor block for a bullet slot.

    docxtpl's {%p ... %} syntax requires the for tag, content, and endfor
    tag to be on SEPARATE paragraphs — putting them on one line is a
    Jinja syntax error.
    """
    doc.add_paragraph("{%p for b in " + slot_id + "_bullets %}")
    doc.add_paragraph("{{b}}", style="List Bullet")
    doc.add_paragraph("{%p endfor %}")


def write_skeleton_docx(path: Path) -> None:
    """Build a minimal example template.docx with the three slot patterns.

    The user is expected to replace this with their polished template,
    keeping the same Jinja slot conventions.
    """
    doc = Document()

    doc.add_heading("YOUR NAME", level=0)
    doc.add_paragraph("email@example.com | linkedin.com/in/you | github.com/you")

    doc.add_heading("Experience", level=1)

    doc.add_heading("Example Role 1 — Example Company", level=2)
    doc.add_paragraph("Jan 2023 – Present")
    add_bullet_loop(doc, "example_role_1")

    doc.add_heading("Example Role 2 — Previous Company", level=2)
    doc.add_paragraph("Mar 2021 – Dec 2022")
    add_bullet_loop(doc, "example_role_2")

    doc.add_heading("Projects", level=1)

    doc.add_heading("Example Project 1", level=2)
    add_bullet_loop(doc, "example_project_1")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Languages: {{skills_languages}}")
    doc.add_paragraph("AI & ML: {{skills_aiml}}")
    doc.add_paragraph("Concepts & Tools: {{skills_concepts_tools}}")

    doc.save(str(path))


def main() -> None:
    targets = [
        (MASTER_MD, MASTER_MD_SKELETON, "text"),
        (TEMPLATE_CONFIG, TEMPLATE_CONFIG_SKELETON, "text"),
        (TEMPLATE_DOCX, None, "docx"),
    ]

    existing = [p for p, _, _ in targets if p.exists()]
    if existing:
        print("Refusing to overwrite existing files:")
        for p in existing:
            print(f"  {p.relative_to(PROJECT_ROOT)}")
        print("\nDelete or move them first if you want a fresh scaffold.")
        sys.exit(1)

    MASTER_MD.write_text(MASTER_MD_SKELETON, encoding="utf-8")
    print(f"Wrote {MASTER_MD.relative_to(PROJECT_ROOT)}")

    TEMPLATE_CONFIG.write_text(TEMPLATE_CONFIG_SKELETON, encoding="utf-8")
    print(f"Wrote {TEMPLATE_CONFIG.relative_to(PROJECT_ROOT)}")

    write_skeleton_docx(TEMPLATE_DOCX)
    print(f"Wrote {TEMPLATE_DOCX.relative_to(PROJECT_ROOT)}")

    print(
        "\nNext steps:\n"
        "  1. Replace template.docx with your polished resume. For each bullet\n"
        "     list, use THREE separate paragraphs (the for/content/endfor must\n"
        "     not be on one line):\n"
        "         {%p for b in <id>_bullets %}\n"
        "         {{b}}                            <-- List Bullet style\n"
        "         {%p endfor %}\n"
        "     For skills, use single placeholders on their lines:\n"
        "         {{skills_languages}} / {{skills_aiml}} / {{skills_concepts_tools}}\n"
        "  2. Edit template-config.yaml so each id matches a slot in your .docx.\n"
        "  3. Fill master.md with your full career narrative + skills inventory.\n"
        "  4. Run /stitch and paste a JD to test the pipeline.\n"
    )


if __name__ == "__main__":
    main()
